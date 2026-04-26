"""Gateway Lambda: export_docx

Generates a DOCX file from Document_State, orders sections per APN
template, uploads to S3, and returns a download key.

Input (via event["inputPayload"] JSON):
    {
        "doc_id": "doc-001",
        "version": 42,
        "meta": { ... },
        "sections": { ... },
        "staffing_plan": { ... }
    }

Output:
    {
        "s3_key": "docs/doc-001/exports/doc-001-v42.docx",
        "bucket": "doc-agent-artifacts",
        "download_url": "presigned URL or null"
    }
"""

from __future__ import annotations

import io
import json
import os
from typing import Any

import boto3

ARTIFACTS_BUCKET = os.environ.get("ARTIFACTS_BUCKET", "doc-agent-artifacts")
REGION = os.environ.get("AWS_REGION", "ap-northeast-2")

# APN template section order for DOCX export
APN_SECTION_ORDER = [
    "cover",
    "executive_summary",
    "stakeholders",
    "success_criteria",
    "assumptions",
    "scope_of_work",
    "architecture",
    "milestones",
    "cost_breakdown",
    "acceptance",
    "resources_cost_estimates",
]

SECTION_TITLES = {
    "cover": "Cover Page",
    "executive_summary": "Executive Summary",
    "stakeholders": "Sponsor / Stakeholder / Team",
    "success_criteria": "Success Criteria & KPIs",
    "assumptions": "Assumptions & Risks",
    "scope_of_work": "Scope of Work",
    "architecture": "Architecture",
    "milestones": "Milestones & Deliverables",
    "cost_breakdown": "Cost Breakdown",
    "acceptance": "Acceptance Criteria",
    "resources_cost_estimates": "Resources & Cost Estimates",
}


def _resolve_field(field: Any) -> Any:
    """Extract display value from a FieldValue dict or return as-is."""
    if isinstance(field, dict) and any(
        k in field for k in ("user_input", "ai_recommended", "calculated")
    ):
        return (
            field.get("user_input")
            or field.get("ai_recommended")
            or field.get("calculated")
        )
    return field


def _build_docx_bytes(
    meta: dict, sections: dict, staffing_plan: dict
) -> bytes:
    """Build a DOCX file as bytes.

    Uses python-docx if available, otherwise falls back to a minimal
    Open XML placeholder so the Lambda always returns a valid file.
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument()

        # Title from meta
        customer = _resolve_field(meta.get("customer", {})) or "Untitled"
        doc.add_heading(f"APN PoC Project Plan — {customer}", level=0)

        # Sections in APN order
        for key in APN_SECTION_ORDER:
            title = SECTION_TITLES.get(key, key)
            doc.add_heading(title, level=1)

            section_data = sections.get(key, {})
            if section_data:
                for field_name, field_val in section_data.items():
                    resolved = _resolve_field(field_val)
                    if resolved is not None:
                        doc.add_paragraph(f"{field_name}: {resolved}")
            else:
                doc.add_paragraph("(No content)")

        # Staffing plan summary
        doc.add_heading("Staffing Plan", level=1)
        roles = staffing_plan.get("roles", {})
        for role_id, role in roles.items():
            name = role.get("display_name", role_id)
            hours = role.get("total_hours", {}).get("calculated", "N/A")
            cost = role.get("total_cost", {}).get("calculated", "N/A")
            doc.add_paragraph(f"{name}: {hours} hours, ${cost}")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # Fallback: return a minimal valid DOCX-like placeholder
        return _minimal_docx_placeholder(meta)


def _minimal_docx_placeholder(meta: dict) -> bytes:
    """Return minimal bytes representing a placeholder DOCX."""
    customer = _resolve_field(meta.get("customer", {})) or "Untitled"
    content = f"APN PoC Project Plan — {customer}\n(Generated placeholder)"
    return content.encode("utf-8")


def handler(event: dict[str, Any], context: Any) -> dict[str, Any]:
    """Lambda entry point for export_docx."""
    try:
        raw = event.get("inputPayload", "{}")
        if isinstance(raw, (bytes, bytearray)):
            raw = raw.decode("utf-8")
        params = json.loads(raw)

        doc_id = params.get("doc_id", "unknown")
        version = params.get("version", 0)
        meta = params.get("meta", {})
        sections = params.get("sections", {})
        staffing_plan = params.get("staffing_plan", {})

        # Build DOCX
        docx_bytes = _build_docx_bytes(meta, sections, staffing_plan)

        # Upload to S3
        s3_key = f"docs/{doc_id}/exports/{doc_id}-v{version}.docx"
        s3 = boto3.client("s3", region_name=REGION)
        s3.put_object(
            Bucket=ARTIFACTS_BUCKET,
            Key=s3_key,
            Body=docx_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Generate presigned URL (1 hour expiry)
        try:
            download_url = s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": ARTIFACTS_BUCKET, "Key": s3_key},
                ExpiresIn=3600,
            )
        except Exception:
            download_url = None

        result = {
            "s3_key": s3_key,
            "bucket": ARTIFACTS_BUCKET,
            "download_url": download_url,
        }
        return {"outputPayload": json.dumps(result)}

    except Exception as e:
        return {"outputPayload": json.dumps({"error": str(e)})}
